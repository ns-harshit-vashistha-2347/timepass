"""
Document Processor for RAG Master Hub.

Handles:
  - Text extraction  : PDF (pdfplumber), DOCX (python-docx), Markdown, plain text
  - Parent-Child chunking : large parent windows for LLM context, small children for retrieval
  - Dense indexing   : sentence-transformers embeddings → ChromaDB (cosine, HNSW)
  - Sparse indexing  : tokenised corpus persisted to disk for BM25 queries
"""

import re
import json
import logging
from functools import lru_cache
from pathlib import Path

import chromadb
from sentence_transformers import SentenceTransformer

from app.core.config import settings

logger = logging.getLogger(__name__)


# ── Singletons (loaded once per worker process) ────────────────────────────

@lru_cache(maxsize=1)
def get_embedding_model() -> SentenceTransformer:
    logger.info(f"[RAG] Loading embedding model: {settings.RAG_EMBEDDING_MODEL}")
    return SentenceTransformer(settings.RAG_EMBEDDING_MODEL)


@lru_cache(maxsize=1)
def get_chroma_client() -> chromadb.PersistentClient:
    chroma_path = Path(settings.RAG_DATA_DIR) / "chroma"
    chroma_path.mkdir(parents=True, exist_ok=True)
    return chromadb.PersistentClient(path=str(chroma_path))


# ── Text Extraction ────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    import io
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"[Page {page_num}]\n{text.strip()}")
    return "\n\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _extract_markdown(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="replace")
    text = re.sub(r"#{1,6}\s+", "", text)           # remove heading markers
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)    # bold → plain
    text = re.sub(r"\*(.+?)\*", r"\1", text)         # italic → plain
    text = re.sub(r"`{1,3}[^`\n]+`{1,3}", " ", text) # inline code
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)      # images
    text = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", text)  # links → label
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to the correct extractor based on file extension."""
    ft = file_type.lower().lstrip(".")
    if ft == "pdf":
        return _extract_pdf(file_bytes)
    elif ft in ("docx", "doc"):
        return _extract_docx(file_bytes)
    elif ft in ("md", "markdown"):
        return _extract_markdown(file_bytes)
    else:  # txt, csv, etc.
        return file_bytes.decode("utf-8", errors="replace")


# ── Sentence-Aware Splitter ────────────────────────────────────────────────

def _split_sentences(text: str) -> list[str]:
    """Split on sentence boundaries, skipping tiny fragments."""
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z\[\(])", text)
    return [p.strip() for p in parts if len(p.strip()) > 15]


def _make_chunks(text: str, chunk_size: int, overlap: int) -> list[str]:
    """
    Build character-bounded chunks while respecting sentence boundaries.
    Carries an `overlap`-character tail into the next chunk for continuity.
    """
    sentences = _split_sentences(text)
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for sent in sentences:
        if current_len + len(sent) > chunk_size and current:
            chunks.append(" ".join(current))
            # Build overlap tail from the end of current chunk
            tail: list[str] = []
            tail_len = 0
            for s in reversed(current):
                if tail_len + len(s) + 1 > overlap:
                    break
                tail.insert(0, s)
                tail_len += len(s) + 1
            current = tail
            current_len = tail_len

        current.append(sent)
        current_len += len(sent) + 1

    if current:
        chunks.append(" ".join(current))

    return [c for c in chunks if c.strip()]


# ── Parent-Child Chunking ──────────────────────────────────────────────────

def chunk_document(
    text: str,
    doc_id: str,
    filename: str,
) -> tuple[list[dict], list[dict]]:
    """
    Two-level chunking:
      parent chunks  → large windows (~1 200 chars) fed to the LLM
      child chunks   → small retrieval units (~350 chars) embedded and searched

    Returns (parent_chunks, child_chunks).
    """
    parent_texts = _make_chunks(
        text,
        chunk_size=settings.RAG_PARENT_CHUNK_SIZE,
        overlap=settings.RAG_PARENT_OVERLAP,
    )

    parents: list[dict] = []
    children: list[dict] = []

    for p_idx, p_text in enumerate(parent_texts):
        parent_id = f"{doc_id}__p{p_idx}"
        parents.append({
            "id":          parent_id,
            "text":        p_text,
            "doc_id":      doc_id,
            "filename":    filename,
            "chunk_index": p_idx,
        })

        child_texts = _make_chunks(
            p_text,
            chunk_size=settings.RAG_CHILD_CHUNK_SIZE,
            overlap=settings.RAG_CHILD_OVERLAP,
        )
        for c_idx, c_text in enumerate(child_texts):
            children.append({
                "id":           f"{parent_id}__c{c_idx}",
                "text":         c_text,
                "parent_id":    parent_id,
                "doc_id":       doc_id,
                "filename":     filename,
                "parent_index": p_idx,
                "child_index":  c_idx,
            })

    return parents, children


# ── Disk Persistence (parents + BM25 corpus) ───────────────────────────────

def _session_dir(session_id: str) -> Path:
    p = Path(settings.RAG_DATA_DIR) / "sessions" / session_id
    p.mkdir(parents=True, exist_ok=True)
    return p


def load_parents(session_id: str) -> dict[str, str]:
    """Return {parent_id: parent_text} from disk."""
    path = _session_dir(session_id) / "parents.json"
    if not path.exists():
        return {}
    with open(path) as f:
        return json.load(f)


def save_parents(session_id: str, new_parents: list[dict]) -> None:
    existing = load_parents(session_id)
    for p in new_parents:
        existing[p["id"]] = p["text"]
    with open(_session_dir(session_id) / "parents.json", "w") as f:
        json.dump(existing, f, ensure_ascii=False)


def load_corpus(session_id: str) -> list[dict]:
    """Return the full BM25 corpus list from disk."""
    path = _session_dir(session_id) / "corpus.json"
    if not path.exists():
        return []
    with open(path) as f:
        return json.load(f)


def save_corpus(session_id: str, new_children: list[dict]) -> None:
    corpus = load_corpus(session_id)
    existing_ids = {c["id"] for c in corpus}
    for c in new_children:
        if c["id"] not in existing_ids:
            corpus.append({
                "id":        c["id"],
                "text":      c["text"],
                "tokens":    c["text"].lower().split(),
                "parent_id": c["parent_id"],
                "doc_id":    c["doc_id"],
                "filename":  c["filename"],
            })
    with open(_session_dir(session_id) / "corpus.json", "w") as f:
        json.dump(corpus, f, ensure_ascii=False)


def delete_doc_from_storage(session_id: str, doc_id: str) -> None:
    """Remove all traces of a document from disk and ChromaDB."""
    # parents.json
    parents = load_parents(session_id)
    parents = {k: v for k, v in parents.items() if not k.startswith(doc_id)}
    with open(_session_dir(session_id) / "parents.json", "w") as f:
        json.dump(parents, f, ensure_ascii=False)

    # corpus.json
    corpus = load_corpus(session_id)
    corpus = [c for c in corpus if c.get("doc_id") != doc_id]
    with open(_session_dir(session_id) / "corpus.json", "w") as f:
        json.dump(corpus, f, ensure_ascii=False)

    # ChromaDB
    try:
        col = get_chroma_client().get_collection(f"rag_{session_id}")
        results = col.get(where={"doc_id": doc_id})
        if results["ids"]:
            col.delete(ids=results["ids"])
    except Exception as e:
        logger.warning(f"[RAG] ChromaDB delete warning: {e}")


# ── Main Indexing Function (CPU-bound — run in executor) ───────────────────

def index_document(
    session_id: str,
    doc_id: str,
    file_bytes: bytes,
    file_type: str,
    filename: str,
) -> dict:
    """
    Full ingestion pipeline:
      extract text → parent-child chunk → embed children → store in ChromaDB + BM25 corpus

    Returns {chunk_count: int, char_count: int}
    Raises ValueError on empty/unreadable documents.
    """
    # 1. Extract
    text = extract_text(file_bytes, file_type)
    if not text.strip():
        raise ValueError("Could not extract any text from the document.")

    # 2. Chunk
    parents, children = chunk_document(text, doc_id, filename)
    if not children:
        raise ValueError("Document produced no indexable chunks after splitting.")

    # 3. Embed children
    model = get_embedding_model()
    child_texts = [c["text"] for c in children]
    embeddings = model.encode(child_texts, batch_size=32, show_progress_bar=False)

    # 4. Store in ChromaDB
    client  = get_chroma_client()
    col_name = f"rag_{session_id}"
    try:
        col = client.get_collection(col_name)
    except Exception:
        col = client.create_collection(
            name=col_name,
            metadata={"hnsw:space": "cosine"},
        )

    col.add(
        ids=[c["id"] for c in children],
        embeddings=embeddings.tolist(),
        documents=child_texts,
        metadatas=[{
            "parent_id": c["parent_id"],
            "doc_id":    c["doc_id"],
            "filename":  c["filename"],
        } for c in children],
    )

    # 5. Persist parents + BM25 corpus to disk
    save_parents(session_id, parents)
    save_corpus(session_id, children)

    logger.info(
        f"[RAG] Indexed doc={doc_id} | parents={len(parents)} | children={len(children)}"
    )
    return {"chunk_count": len(children), "char_count": len(text)}